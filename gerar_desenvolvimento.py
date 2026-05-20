"""
Gera a Seção 4 - Desenvolvimento no TCC_Formatado.docx,
inserindo texto acadêmico, tabelas e gráficos gerados a partir
das bases de dados reais do projeto.
"""

import os, io, warnings
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.chdir(r'c:\tcc-uvv-2026')

# ══════════════════════════════════════════════════════════════
# 1. CARREGAMENTO E PRÉ-PROCESSAMENTO DOS DADOS
# ══════════════════════════════════════════════════════════════
print("Carregando bases...")
df_base = pd.read_csv('Bases/base_unificada_2021_2026.csv', sep=';', parse_dates=['data'])
df_prev = pd.read_csv('Bases/previsao_volume_2026.csv', parse_dates=['data'])
print(f"  base_unificada : {len(df_base):,} linhas")

MINUTOS_POR_COLAB = 480

df_cap = (
    df_base
    .assign(min_demanda=lambda d: d['volume_atendimentos'] * d['tempo_medio_min'])
    .groupby(['data', 'empresa', 'tipo'], as_index=False)
    .agg(
        volume_total =('volume_atendimentos', 'sum'),
        demanda_min  =('min_demanda',         'sum'),
        colaboradores=('quantidade_colaboradores', 'first'),
    )
)
df_cap['capacidade_min'] = df_cap['colaboradores'] * MINUTOS_POR_COLAB
df_cap['utilizacao_pct'] = (df_cap['demanda_min'] / df_cap['capacidade_min'] * 100).round(1)
df_cap['gap_min']        = (df_cap['capacidade_min'] - df_cap['demanda_min']).round(1)
df_cap['ano']            = df_cap['data'].dt.year
df_cap['status']         = pd.cut(
    df_cap['utilizacao_pct'], bins=[-np.inf, 70, 90, np.inf],
    labels=['Adequado', 'Alerta', 'Critico']
).astype(str)

df_prev_2026 = df_prev[(df_prev['data'] >= '2026-01-01') & (df_prev['data'] <= '2026-12-31')].copy()
df_prev_2026['mes'] = df_prev_2026['data'].dt.to_period('M')

# ══════════════════════════════════════════════════════════════
# 2. GERAÇÃO DOS GRÁFICOS
# ══════════════════════════════════════════════════════════════
print("Gerando gráficos...")

plt.rcParams.update({'font.size': 11, 'axes.titlesize': 12, 'axes.labelsize': 11})

def savefig(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0); plt.close(fig)
    return buf

# G1 – Curva sazonal
ctrl = np.array([
    (0,1.30),(45,1.25),(80,1.05),(105,0.88),(135,0.80),
    (182,0.78),(228,0.78),(260,0.85),(290,1.05),(320,1.25),
    (355,1.40),(366,1.30)
])
dias = np.arange(0, 366)
fat  = np.interp(dias, ctrl[:,0], ctrl[:,1])
fig, ax = plt.subplots(figsize=(10, 4))
ax.plot(dias, fat, color='steelblue', lw=2.5)
ax.fill_between(dias, 1.0, fat, where=fat>=1.0, alpha=0.18, color='tomato',    label='Acima da média')
ax.fill_between(dias, fat, 1.0, where=fat< 1.0, alpha=0.18, color='goldenrod', label='Abaixo da média')
ax.axhline(1.0, color='gray', ls='--', lw=0.8)
ax.set_xticks([1,32,60,91,121,152,182,213,244,274,305,335])
ax.set_xticklabels(['Jan','Fev','Mar','Abr','Mai','Jun','Jul','Ago','Set','Out','Nov','Dez'])
ax.set_xlabel('Mês'); ax.set_ylabel('Fator Sazonal')
ax.set_title('Curva de Sazonalidade Intra-Anual')
ax.legend(loc='lower center'); plt.tight_layout()
img_sazonal = savefig(fig)

# G2 – Volume histórico diário
serie = (df_base[df_base['tipo']=='Historico']
         .groupby('data')['volume_atendimentos'].sum())
fig, ax = plt.subplots(figsize=(12, 4))
serie.plot(ax=ax, color='steelblue', lw=0.7, alpha=0.9)
ax.set_title('Volume Diário Total de Atendimentos — Histórico 2021–2025')
ax.set_ylabel('Volume de Atendimentos')
ax.xaxis.set_major_formatter(mdates.DateFormatter('%b/%Y'))
plt.tight_layout()
img_vol_hist = savefig(fig)

# G3 – Comparação dos modelos (mensal 2026)
mensal = df_prev_2026.groupby('mes')[['ARIMA','Prophet','RandomForest']].sum()
mensal.index = ['Jan','Fev','Mar','Abr','Mai','Jun','Jul','Ago','Set','Out','Nov','Dez']
fig, ax = plt.subplots(figsize=(12, 5))
ax.plot(mensal.index, mensal['ARIMA'],        color='tomato',     marker='o', ms=5, lw=2, label='ARIMA')
ax.plot(mensal.index, mensal['Prophet'],      color='darkorange', marker='s', ms=5, lw=2, label='Prophet')
ax.plot(mensal.index, mensal['RandomForest'], color='seagreen',   marker='^', ms=5, lw=2, label='Random Forest')
ax.set_title('Comparação dos Modelos — Volume Mensal Previsto 2026')
ax.set_ylabel('Volume de Atendimentos'); ax.grid(axis='y', alpha=0.3)
ax.legend(); plt.tight_layout()
img_forecast = savefig(fig)

total_arima   = int(df_prev_2026['ARIMA'].sum())
total_prophet = int(df_prev_2026['Prophet'].sum())
total_rf      = int(df_prev_2026['RandomForest'].sum())

# G4 – Status por ano (stacked bar)
status_ano = (df_cap.groupby(['ano','status']).size().unstack(fill_value=0))
for c in ['Adequado','Alerta','Critico']:
    if c not in status_ano.columns: status_ano[c] = 0
status_ano = status_ano[['Adequado','Alerta','Critico']]
fig, ax = plt.subplots(figsize=(10, 5))
status_ano.plot(kind='bar', stacked=True, ax=ax,
    color={'Adequado':'seagreen','Alerta':'goldenrod','Critico':'tomato'})
ax.set_title('Distribuição do Status de Capacidade por Ano (empresa × dia)')
ax.set_ylabel('Ocorrências'); ax.set_xlabel('')
ax.set_xticklabels(status_ano.index.astype(str), rotation=0)
ax.legend(title='Status'); plt.tight_layout()
img_status_ano = savefig(fig)

# G5 – Top 10 críticas 2026
top_crit = (df_cap[df_cap['ano']==2026]
            .groupby('empresa')['utilizacao_pct'].mean()
            .sort_values(ascending=False).head(10))
fig, ax = plt.subplots(figsize=(10, 5))
top_crit.plot(kind='bar', color='tomato', ax=ax)
ax.axhline(90, color='darkred', ls='--', lw=1.2, label='Limite Crítico (90%)')
ax.axhline(100, color='black',  ls=':',  lw=1.0, label='100%')
ax.set_title('Top 10 Empresas — Maior Utilização Média Prevista (2026)')
ax.set_ylabel('Utilização Média (%)'); ax.set_xlabel('')
ax.set_xticklabels(top_crit.index, rotation=45, ha='right')
ax.legend(); plt.tight_layout()
img_top_crit = savefig(fig)

# G6 – Top 10 ociosas 2026
top_ocio = (df_cap[df_cap['ano']==2026]
            .groupby('empresa')['utilizacao_pct'].mean()
            .sort_values().head(10))
fig, ax = plt.subplots(figsize=(10, 5))
top_ocio.plot(kind='bar', color='steelblue', ax=ax)
ax.axhline(70, color='goldenrod', ls='--', lw=1.2, label='Limite Adequado (70%)')
ax.set_title('Top 10 Empresas — Menor Utilização Média Prevista (2026)')
ax.set_ylabel('Utilização Média (%)'); ax.set_xlabel('')
ax.set_xticklabels(top_ocio.index, rotation=45, ha='right')
ax.legend(); plt.tight_layout()
img_top_ocio = savefig(fig)

# ── Métricas para o texto ─────────────────────────────────────
df_mensal_2026 = (
    df_cap[df_cap['ano']==2026]
    .assign(mes=lambda d: d['data'].dt.to_period('M').astype(str))
    .groupby('mes')[['utilizacao_pct','gap_min','volume_total']].mean().round(1)
)
status_2026 = df_cap[df_cap['ano']==2026]['status'].value_counts()
n_criticas  = int(status_2026.get('Critico', 0))
n_adequadas = int(status_2026.get('Adequado', 0))
n_alerta    = int(status_2026.get('Alerta',   0))
util_media_hist = df_cap[df_cap['tipo']=='Historico']['utilizacao_pct'].mean()
util_media_2026 = df_cap[df_cap['ano']==2026]['utilizacao_pct'].mean()

# Colab sugerido
df_colab_2026 = (
    df_cap[df_cap['ano']==2026]
    .groupby('empresa', as_index=False)
    .agg(util_media=('utilizacao_pct','mean'),
         demanda_media=('demanda_min','mean'),
         colaboradores=('colaboradores','first'))
)
df_colab_2026['colab_sug'] = np.ceil(df_colab_2026['demanda_media']/MINUTOS_POR_COLAB).astype(int)
df_colab_2026['diferenca'] = df_colab_2026['colab_sug'] - df_colab_2026['colaboradores']
criticas_info  = df_colab_2026.sort_values('util_media', ascending=False).head(10)
ociosas_info   = df_colab_2026.sort_values('util_media').head(10)

print("Gráficos prontos. Editando documento...")

# ══════════════════════════════════════════════════════════════
# 3. EDIÇÃO DO DOCX
# ══════════════════════════════════════════════════════════════
doc = Document('TCC_Formatado.docx')

# Localiza o parágrafo "Referências"
ref_el = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Refer') and 'Heading' in p.style.name:
        ref_el = p._element; break
assert ref_el is not None, "Seção Referências não encontrada."

# ── Helpers ────────────────────────────────────────────────────

def ins(style='normal'):
    p = doc.add_paragraph(style=style)
    ref_el.addprevious(p._element)
    return p

def h1(text):
    p = ins('Heading 1'); p.add_run(text); return p

def h2(text):
    p = ins('Heading 2'); p.add_run(text); return p

def h3(text):
    p = ins('Heading 3'); p.add_run(text); return p

def txt(*parts):
    """parts = str  |  (str, bold, italic)"""
    p = ins('normal')
    for x in parts:
        if isinstance(x, str):
            p.add_run(x)
        else:
            r = p.add_run(x[0])
            if len(x) > 1: r.bold   = x[1]
            if len(x) > 2: r.italic = x[2]
    return p

def B(t):  return (t, True,  False)
def I(t):  return (t, False, True)
def BI(t): return (t, True,  True)

def img(buf, caption=None, w=5.5):
    p = ins('normal'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(w))
    if caption:
        c = ins('normal'); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(10)

def shade_cell(cell, hex_color='2E75B6'):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pr.append(shd)

def set_table_borders(t):
    """Add visible borders to a table via XML."""
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def table(headers, rows, caption=None):
    if caption:
        p = ins('normal'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.bold = True; r.font.size = Pt(10)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'TableNormal'
    set_table_borders(t)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            t.rows[ri+1].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_el.addprevious(t._tbl)
    return t

def blank():
    ins('normal')

# ══════════════════════════════════════════════════════════════
# 4. CONTEÚDO — SEÇÃO 4: DESENVOLVIMENTO
# ══════════════════════════════════════════════════════════════

h1('4 Desenvolvimento')

# ──────────────────────────────────────────
h2('4.1 Caracterização do Cenário de Análise')
# ──────────────────────────────────────────

txt('O presente estudo foi conduzido sobre um conjunto de dados simulados que representa o ambiente operacional de uma rede de empresas prestadoras de serviços no setor de reparação e instalação de vidros automotivos. O cenário contempla ', B('80 empresas clientes'), ' (denominadas EMP1 a EMP80), cada uma demandando serviços em cinco categorias de produtos: Parabrisa, Vigia, Laterais, Farol e Retrovisor.')

txt('A série histórica abrange o período de ', B('1º de janeiro de 2021 a 31 de dezembro de 2025'), ' (1.826 dias), totalizando 730.400 registros de volume de atendimentos. Para cada combinação empresa–produto, foram registrados diariamente: o volume de atendimentos realizados, o tempo médio de execução por atendimento (em minutos) e a quantidade de colaboradores disponíveis na empresa.')

txt('A escolha por dados simulados justifica-se pela necessidade de construir um ambiente controlado, onde os padrões de sazonalidade, crescimento e variabilidade fossem conhecidos a priori, permitindo avaliar a eficácia dos modelos de previsão e a consistência dos indicadores de capacidade produtiva gerados. A simulação foi implementada integralmente em Python, utilizando as bibliotecas Pandas e NumPy, e os parâmetros foram fixados com semente aleatória (seed 42) para garantir reprodutibilidade.')

# ──────────────────────────────────────────
h2('4.2 Construção da Base de Dados')
# ──────────────────────────────────────────

h3('4.2.1 Base de Volume de Atendimentos')

txt('A base de volume de atendimentos foi gerada com base em três componentes principais: sazonalidade intra-anual, tendência de crescimento e variabilidade estocástica. A combinação dessas três forças resulta em uma série temporal realista, com padrões cíclicos e ruído natural.')

txt(B('Sazonalidade intra-anual:'), ' o comportamento sazonal foi modelado por interpolação linear entre pontos de controle distribuídos ao longo do ano, associando cada dia do ano a um fator multiplicador. O resultado é uma curva suave que reproduz o padrão típico do setor automotivo: picos no início (janeiro) e no final do ano (dezembro), e vale no período de julho a agosto, conforme ilustrado na Figura 1 e detalhado na Tabela 1.')

blank()
img(img_sazonal, 'Figura 1 – Curva de sazonalidade intra-anual aplicada à geração do volume de atendimentos', w=5.5)
blank()

table(
    ['Dia do Ano', 'Mês Aproximado', 'Fator Sazonal'],
    [('0','01/jan','1,30'),('45','meados fev','1,25'),('80','meados mar','1,05'),
     ('105','meados abr','0,88'),('135','meados mai','0,80'),('182','meados jul','0,78'),
     ('228','meados ago','0,78'),('260','meados set','0,85'),('290','meados out','1,05'),
     ('320','meados nov','1,25'),('355','meados dez','1,40'),('366','31/dez','1,30')],
    caption='Tabela 1 – Pontos de controle da sazonalidade intra-anual'
)
blank()

txt(B('Tendência de crescimento:'), ' foi aplicado um fator de crescimento composto de 5% ao ano a partir de 2021, calculado como ', BI('crescimento = 1,05^(ano − 2021)'), '. Esse fator reflete a expansão gradual do volume de serviços ao longo do horizonte histórico, resultando em aproximadamente 21,6% de crescimento acumulado entre 2021 e 2025.')

txt(B('Variabilidade estocástica:'), ' o volume diário por empresa/produto foi amostrado de uma distribuição de Poisson com média igual ao produto do volume base, fator sazonal e fator de crescimento. A distribuição de Poisson é especialmente adequada para modelar contagens de eventos discretos — como o número de atendimentos diários —, pois captura naturalmente a variabilidade intrínseca de processos de chegada aleatória.')

txt('Os volumes base por produto foram definidos conforme a Tabela 2, refletindo a proporção típica de demanda por tipo de serviço no setor de vidros automotivos.')

table(
    ['Produto', 'Volume Base (atend./empresa/dia)'],
    [('Parabrisa','30'),('Vigia','25'),('Laterais','22'),('Farol','15'),('Retrovisor','10')],
    caption='Tabela 2 – Volume base de atendimentos por produto'
)
blank()

txt('A Figura 2 apresenta a série de volume diário total (agregado sobre as 80 empresas e 5 produtos) para o período histórico de 2021 a 2025, onde são visíveis os padrões sazonal e de crescimento incorporados na geração dos dados.')

blank()
img(img_vol_hist, 'Figura 2 – Volume diário total de atendimentos — histórico 2021–2025', w=6.0)
blank()

h3('4.2.2 Base de Tempo Médio de Atendimento')

txt('A base de tempo médio por atendimento foi construída para refletir a heterogeneidade operacional entre empresas e a variabilidade natural do processo de execução. Cada produto possui um ', B('tempo base de referência'), ', conforme a Tabela 3, representando a duração média esperada em condições normais de trabalho.')

table(
    ['Produto', 'Tempo Médio Base (min)'],
    [('Parabrisa','45'),('Vigia','65'),('Laterais','55'),('Farol','35'),('Retrovisor','30')],
    caption='Tabela 3 – Tempo médio base por produto'
)
blank()

txt('Para capturar a diferença de eficiência entre empresas, foi atribuído a cada uma delas um ', B('fator de velocidade'), ' amostrado de uma distribuição uniforme entre 0,85 e 1,15 (variação de ±15% em relação ao tempo base). Empresas com fator inferior a 1,0 executam os serviços mais rapidamente; empresas com fator superior operam com maior tempo de execução.')

txt('A variação diária foi modelada por uma distribuição normal com média igual ao tempo esperado (tempo base × fator da empresa) e desvio padrão de 7 minutos, com piso de 10 minutos para preservar a plausibilidade física dos registros. Esse desvio padrão reproduz variações naturais decorrentes de fatores como complexidade específica do serviço, disponibilidade de equipamentos e condições do veículo atendido.')

txt('Segundo Barnes (1989), o estudo de tempos e movimentos é fundamental para o dimensionamento correto da capacidade produtiva. No contexto deste trabalho, a definição de tempos padrão por produto — com variação controlada por empresa — permite calcular, com precisão, a carga de trabalho imposta a cada equipe, viabilizando a análise de utilização de capacidade descrita nas seções subsequentes.')

h3('4.2.3 Dimensionamento de Colaboradores por Dia')

txt('A base de colaboradores por dia foi gerada simulando três perfis reais de dimensionamento de equipes, derivados da carga operacional histórica de cada empresa. A premissa central é que cada colaborador dispõe de ', B('480 minutos úteis por dia'), ' (jornada de 8 horas), conforme referencial de capacidade adotado em Slack, Chambers e Johnston (2009).')

txt('O número de ', B('colaboradores ideal'), ' foi calculado como a razão entre a carga média histórica e os 480 minutos disponíveis por colaborador. A partir desse ideal, cada empresa recebeu um perfil de dimensionamento, sorteado com semente 42 para reprodutibilidade, conforme a Tabela 4:')

table(
    ['Perfil', '% Empresas', 'Fator s/ Ideal', 'Utilização Média Esperada'],
    [('Adequado','~30%','1,60×','~62% — folga de pessoal'),
     ('Alerta',  '~40%','1,20×','~83% — quadro saudável'),
     ('Crítico', '~30%','0,85×','~118% — sobrecarga operacional')],
    caption='Tabela 4 – Perfis de dimensionamento de colaboradores'
)
blank()

txt('O perfil ', B('"Adequado"'), ' representa empresas com excesso de pessoal — cenário comum em organizações que contratam preventivamente ou que enfrentaram redução de demanda sem ajuste do quadro. O perfil ', B('"Crítico"'), ' simula empresas com subcontratação crônica, operando sistematicamente acima da capacidade disponível, o que gera riscos de qualidade e esgotamento das equipes.')

txt('Adicionalmente, uma variação diária de ±1 colaborador foi aplicada nos dias em que o volume de atendimentos desviou mais de 15% da média histórica da empresa, simulando os ajustes de escala (cobertura por outros colaboradores, horas extras, afastamentos) típicos de ambientes produtivos reais.')

# ──────────────────────────────────────────
h2('4.3 Modelos de Previsão de Demanda para 2026')
# ──────────────────────────────────────────

txt('Para a projeção do volume de atendimentos em 2026, foram desenvolvidos e comparados três modelos de previsão de séries temporais: ARIMA/SARIMA, Prophet e Random Forest. Os modelos foram aplicados sobre a série de ', B('volume total diário'), ' — soma de todas as empresas e produtos —, abrangendo 1.826 observações de 2021 a 2025. Essa abordagem de previsão agregada seguida de desagregação proporcional é recomendada quando o horizonte de previsão disponível por segmento é insuficiente para calibrar modelos individuais (MAKRIDAKIS; WHEELWRIGHT; HYNDMAN, 1998).')

h3('4.3.1 Preparação da Série Temporal')

txt('A série temporal foi construída pela agregação diária do volume total de atendimentos de todas as empresas e produtos, resultando em uma sequência contínua sem valores ausentes. A análise preliminar identificou três componentes estruturais relevantes para a modelagem:')

txt('(1) ', B('Sazonalidade semanal:'), ' volumes reduzidos nos fins de semana e superiores nos dias úteis, reflexo do perfil operacional das empresas clientes;')

txt('(2) ', B('Sazonalidade anual:'), ' padrão repetitivo com picos em janeiro e dezembro e vale em julho–agosto, coerente com a curva sazonal incorporada na geração dos dados;')

txt('(3) ', B('Tendência de crescimento:'), ' elevação gradual e consistente do nível da série ao longo dos cinco anos, reflexo do crescimento anual de 5% embutido na simulação.')

txt('O teste de Dickey-Fuller Aumentado (ADF) foi aplicado para verificar a estacionariedade da série. O resultado confirmou que a série apresenta raiz unitária em nível, tornando-se estacionária após uma diferenciação, validando a aplicabilidade do modelo ARIMA com parâmetro d = 1.')

h3('4.3.2 Modelo ARIMA/SARIMA')

txt('O modelo ARIMA — sigla para ', I('Autoregressive Integrated Moving Average'), ' — é amplamente utilizado para previsão de séries temporais univariadas. Sua formulação combina três componentes: a parte autorregressiva (AR), que relaciona o valor presente com observações passadas; a integração (I), que trata a não-estacionariedade por meio de diferenciação; e a média móvel (MA), que incorpora os erros de previsão passados como variáveis explicativas.')

txt('A extensão sazonal, denominada SARIMA, adiciona parâmetros sazonais (P, D, Q) e o período sazonal ', I('m'), ', tornando o modelo apto a capturar padrões periódicos recorrentes. No presente estudo, foi adotada a sazonalidade semanal (', I('m = 7'), '), dada a variação regular entre dias úteis e fins de semana.')

txt('A seleção automática dos parâmetros foi realizada pelo método ', I('auto_arima'), ', que avalia sistematicamente combinações dos parâmetros minimizando o Critério de Informação de Akaike (AIC). Os limites de busca foram: p, q ≤ 3 e P, Q ≤ 2. O método de busca ', I('stepwise'), ' foi empregado para reduzir o tempo computacional sem comprometer a qualidade da seleção.')

txt('O modelo ARIMA/SARIMA apresenta como principais vantagens a interpretabilidade estatística e a robustez para séries com estruturas de autocorrelação bem definidas. Como limitação, a captura simultânea de sazonalidade semanal e anual exige especificações mais complexas, pois o SARIMA com um único período sazonal não modela múltiplas frequências nativamente.')

h3('4.3.3 Modelo Prophet')

txt('O Prophet é um modelo de previsão de séries temporais desenvolvido pela equipe de ciência de dados do Meta (antigo Facebook), projetado para dados de negócios com múltiplas sazonalidades, feriados e tendências não lineares (TAYLOR; LETHAM, 2018). Sua formulação decompõe a série em componentes aditivos ou multiplicativos: tendência, sazonalidade e efeitos de eventos especiais.')

txt('No presente estudo, o Prophet foi configurado com os seguintes parâmetros:')
txt('• ', B('Sazonalidade anual e semanal:'), ' ambas ativadas para capturar os dois ciclos identificados na análise preliminar;')
txt('• ', B('Modo multiplicativo:'), ' mais adequado quando a amplitude das oscilações sazonais cresce proporcionalmente com o nível da série — comportamento observado neste conjunto de dados, onde o crescimento anual amplifica os picos e vales sazonais;')
txt('• ', B('Feriados brasileiros:'), ' incorporados via ', I('make_holidays_df'), ', contemplando datas como Carnaval, Semana Santa, Dia do Trabalho, Dia da Independência e Natal, que afetam sistematicamente o volume de atendimentos;')
txt('• ', B('changepoint_prior_scale = 0,10:'), ' parâmetro que controla a flexibilidade da tendência — valor menor resulta em tendência mais suave, reduzindo o risco de overfitting a variações de curto prazo.')

txt('O Prophet se destaca pela facilidade de incorporação de conhecimento de negócio (como feriados e eventos sazonais customizados), robustez a dados ausentes e boa generalização para horizontes de médio prazo. Essas características, aliadas à consistência das previsões com o padrão histórico observado, justificaram sua seleção como ', B('modelo de referência para a desagregação da demanda 2026'), ' por empresa e produto.')

h3('4.3.4 Modelo Random Forest para Séries Temporais')

txt('O Random Forest é um método de aprendizado de máquina baseado em conjuntos (', I('ensembles'), ') de árvores de decisão. Cada árvore é treinada em uma amostra bootstrap do conjunto de treinamento, e as previsões finais são obtidas pela média das previsões individuais, reduzindo a variância e melhorando a generalização em relação a uma única árvore (BREIMAN, 2001).')

txt('A aplicação a séries temporais requer a transformação da dependência temporal em variáveis explicativas (', I('features'), '), convertendo o problema sequencial em uma tarefa de regressão supervisionada. As seguintes features foram construídas:')
txt('• ', B('Lags (defasagens):'), ' valores da série nos instantes t−1, t−2, t−3, t−7, t−14, t−21, t−28 e t−30, capturando dependências de curto, médio e longo prazo;')
txt('• ', B('Estatísticas rolantes:'), ' média e desvio padrão dos últimos 7, 14 e 30 dias, calculadas com base no dia anterior para evitar vazamento de dados (', I('data leakage'), ');')
txt('• ', B('Features temporais:'), ' dia do ano, dia da semana, semana do ano, mês, trimestre, ano e indicador binário de fim de semana;')
txt('• ', B('Codificação cíclica:'), ' seno e cosseno do dia da semana e do dia do ano, preservando a continuidade matemática dos ciclos e evitando que o modelo interprete a passagem de domingo para segunda-feira como descontinuidade.')

txt('O modelo foi treinado com os dados de 2021 a 2024 e validado no ano de 2025, com 500 estimadores, profundidade máxima de 10 e mínimo de 5 amostras por folha. A previsão para 2026 foi realizada de forma ', B('iterativa'), ': a cada dia previsto, o valor estimado é incorporado ao histórico para alimentar os lags dos dias seguintes, mantendo a consistência temporal da cadeia de previsões.')

h3('4.3.5 Comparação dos Modelos de Previsão')

txt(f'Os três modelos foram comparados quanto ao volume total anual previsto para 2026 e ao comportamento mensal das previsões. A Tabela 5 apresenta os totais anuais estimados por cada abordagem, e a Figura 3 ilustra a comparação mês a mês.')

table(
    ['Modelo', 'Volume Total Previsto 2026', 'Observações'],
    [('ARIMA/SARIMA', f'{total_arima:,}'.replace(',','.'), 'Captura bem a sazonalidade semanal; limitação na sazonalidade anual'),
     ('Prophet',      f'{total_prophet:,}'.replace(',','.'), 'Sazonalidade múltipla + feriados; modelo selecionado para desagregação'),
     ('Random Forest',f'{total_rf:,}'.replace(',','.'), 'Boa validação em 2025; tendência à regressão à média no longo prazo')],
    caption='Tabela 5 – Totais anuais previstos para 2026 por modelo'
)
blank()

blank()
img(img_forecast, 'Figura 3 – Comparação dos modelos de previsão — volume mensal previsto 2026', w=6.0)
blank()

txt('A análise visual revelou que o ', B('Prophet'), ' produziu previsões mais suaves e consistentes com o padrão histórico, especialmente nos meses de pico (janeiro e dezembro). O ARIMA capturou bem a sazonalidade semanal, mas apresentou maior oscilação no horizonte anual. O Random Forest tendeu a regredir à média histórica em horizontes mais distantes — comportamento esperado em métodos baseados em árvores quando aplicados além da faixa de treinamento.')

txt('Considerando a consistência sazonal, a incorporação de feriados e a robustez para previsões de médio prazo, o ', B('Prophet foi selecionado como modelo de referência'), ' para a etapa de análise de capacidade, onde a previsão diária de 2026 é desagregada por empresa e produto com base nas proporções históricas de 2025.')

# ──────────────────────────────────────────
h2('4.4 Análise de Capacidade Produtiva')
# ──────────────────────────────────────────

h3('4.4.1 Consolidação da Base Unificada (2021–2026)')

txt('A base unificada foi construída pela concatenação da série histórica (2021–2025) com a projeção para 2026, totalizando seis anos de análise. Para a projeção, o Prophet fornece o ', B('volume total diário'), ' — soma sobre todas as empresas e produtos —, exigindo uma etapa de desagregação para viabilizar a análise por empresa.')

txt('A desagregação foi realizada com base nas ', B('proporções históricas de 2025'), ': para cada combinação (empresa, produto), calculou-se a participação percentual no volume total do ano. Essas proporções foram multiplicadas pelo volume diário previsto, redistribuindo a demanda de 2026 com a mesma estrutura relativa do último ano histórico. Essa abordagem pressupõe estabilidade na composição da demanda — hipótese razoável para horizontes de previsão de um ano.')

txt('Os demais parâmetros da projeção foram definidos como: (a) ', B('tempo médio de atendimento 2026'), ' = média do tempo observado em 2025, por empresa/produto; (b) ', B('quantidade de colaboradores 2026'), ' = média arredondada do quadro de 2025, por empresa, representando o cenário de manutenção do quadro atual sem reforço ou redução.')

h3('4.4.2 Indicadores de Capacidade Produtiva')

txt('Para cada empresa e dia do período analisado, foram calculados quatro indicadores de capacidade, derivados da premissa de 480 minutos úteis disponíveis por colaborador por dia (SLACK; CHAMBERS; JOHNSTON, 2009), conforme a Tabela 6:')

table(
    ['Indicador', 'Definição', 'Fórmula'],
    [('Demanda (min)',   'Total de minutos necessários para executar todos os atendimentos do dia', 'Σ(volume × tempo_médio) por produto'),
     ('Capacidade (min)','Total de minutos produtivos disponíveis na empresa',                     'colaboradores × 480'),
     ('Utilização (%)',  'Proporção da capacidade efetivamente consumida pela demanda',             '(demanda / capacidade) × 100'),
     ('Gap (min)',       'Saldo entre capacidade e demanda (positivo = folga; negativo = sobrecarga)', 'capacidade − demanda')],
    caption='Tabela 6 – Indicadores de capacidade produtiva calculados por empresa/dia'
)
blank()

txt('A classificação por status operacional foi definida conforme a Tabela 7, com base nos limiares de utilização recomendados para ambientes de serviço:')

table(
    ['Status', 'Critério de Utilização', 'Interpretação Gerencial'],
    [('Adequado','< 70%', 'Folga de capacidade; possível excesso de pessoal'),
     ('Alerta',  '70% a 90%', 'Faixa operacional saudável; baixo risco de ruptura'),
     ('Crítico', '> 90%', 'Risco elevado de sobrecarga, atrasos e degradação do serviço')],
    caption='Tabela 7 – Classificação do status operacional por faixa de utilização'
)
blank()

h3('4.4.3 Resultados Históricos da Capacidade (2021–2025)')

txt(f'A análise da taxa de utilização ao longo do período histórico reflete diretamente os perfis de dimensionamento definidos na geração da base de colaboradores. A taxa de utilização média histórica observada foi de ', B(f'{util_media_hist:.1f}%'), ', coerente com a distribuição de perfis: ~30% das empresas em condição crítica (fator 0,85×), ~40% em alerta (fator 1,2×) e ~30% em condição adequada (fator 1,6×).')

txt('A Figura 4 apresenta a distribuição anual dos status operacionais, evidenciando a persistência dos três grupos ao longo de todos os anos e o leve aumento das ocorrências críticas nas colunas mais recentes, reflexo do crescimento do volume de atendimentos sem ajuste proporcional do quadro de colaboradores.')

blank()
img(img_status_ano, 'Figura 4 – Distribuição do status de capacidade por ano (empresa × dia)', w=5.5)
blank()

txt('A sazonalidade intra-anual é também perceptível na série de utilização diária: os meses de janeiro e dezembro concentram os maiores picos de utilização, enquanto julho e agosto correspondem ao período de menor pressão sobre a capacidade. Esse padrão tem implicações diretas no planejamento de escalas, férias e contratações temporárias.')

h3('4.4.4 Projeção de Capacidade para 2026')

txt(f'A projeção para 2026, com base nas previsões do Prophet, indica uma taxa de utilização média de ', B(f'{util_media_2026:.1f}%'), ' — superior à média histórica de ', B(f'{util_media_hist:.1f}%'), ' —, reflexo do crescimento do volume de atendimentos previsto sem correspondente ampliação do quadro de colaboradores. O cenário simulado representa uma pressão crescente sobre o sistema produtivo, especialmente para as empresas que já operam próximas ao limite.')

txt('A Tabela 8 apresenta o resumo mensal dos indicadores médios de capacidade para 2026, calculados como médias entre todas as empresas:')

mensal_rows = []
for mes, row in df_mensal_2026.iterrows():
    mensal_rows.append((str(mes), f"{row['utilizacao_pct']:.1f}%",
                        f"{row['gap_min']:.0f}", f"{row['volume_total']:.0f}"))

table(
    ['Mês', 'Utilização Média (%)', 'Gap Médio (min/dia)', 'Volume Médio (atend./empresa/dia)'],
    mensal_rows,
    caption='Tabela 8 – Resumo mensal da capacidade produtiva prevista para 2026 (médias entre empresas)'
)
blank()

h3('4.4.5 Identificação de Empresas em Situação Crítica e em Ociosidade')

n_emp_crit = int((df_colab_2026['util_media'] > 90).sum())
n_emp_ocio = int((df_colab_2026['util_media'] < 70).sum())

txt(f'A granularidade da análise por empresa permite identificar dois grupos de interesse estratégico. Das 80 empresas analisadas, ', B(f'{n_emp_crit} operam em situação crítica'), f' (utilização média > 90% ao longo de 2026) e ', B(f'{n_emp_ocio} apresentam ociosidade significativa'), f' (utilização média < 70%).')

txt(B('Empresas em situação crítica'), f' (utilização média > 90% em 2026): essas empresas operam sistematicamente com demanda superior à capacidade disponível, acumulando um gap negativo médio diário que indica quantos minutos de trabalho não são cobertos pelo quadro atual. A Figura 5 apresenta as 10 empresas com maior utilização média prevista.')

blank()
img(img_top_crit, 'Figura 5 – Top 10 empresas com maior utilização média prevista (2026)', w=5.5)
blank()

txt('Para essas empresas, o indicador ', B('colaboradores sugeridos'), ' — calculado como ⌈demanda_média_diária / 480⌉ — quantifica o reforço mínimo necessário. A Tabela 9 apresenta o detalhamento das 10 empresas mais críticas:')

crit_rows = []
for _, row in criticas_info.iterrows():
    diff = int(row['colab_sug'] - row['colaboradores'])
    diff_str = f"+{diff}" if diff > 0 else str(diff)
    crit_rows.append((str(row['empresa']), f"{row['util_media']:.1f}%",
                      str(int(row['colaboradores'])), str(int(row['colab_sug'])), diff_str))
table(
    ['Empresa', 'Utilização Média (%)', 'Colab. Atuais', 'Colab. Sugeridos', 'Diferença'],
    crit_rows,
    caption='Tabela 9 – Top 10 empresas em situação crítica e necessidade de reforço de quadro (2026)'
)
blank()

txt('Em um cenário de sobrecarga persistente, as consequências operacionais incluem: aumento do tempo médio de atendimento por fadiga e pressão de tempo, elevação da taxa de erros e retrabalho, e degradação da satisfação do cliente. A identificação prévia dessas empresas permite uma ação proativa de contratação ou redistribuição de equipes antes que o problema se materialize.')

txt(B('Empresas em ociosidade'), f' (utilização média < 70% em 2026): essas empresas mantêm um quadro de colaboradores significativamente superior ao exigido pela demanda prevista. A Figura 6 apresenta as 10 empresas com menor utilização média.')

blank()
img(img_top_ocio, 'Figura 6 – Top 10 empresas com menor utilização média prevista (2026)', w=5.5)
blank()

txt('A Tabela 10 detalha o potencial de otimização para as 10 empresas mais ociosas:')

ocio_rows = []
for _, row in ociosas_info.iterrows():
    diff = int(row['colab_sug'] - row['colaboradores'])
    diff_str = f"{diff}" if diff >= 0 else str(diff)
    ocio_rows.append((str(row['empresa']), f"{row['util_media']:.1f}%",
                      str(int(row['colaboradores'])), str(int(row['colab_sug'])), diff_str))
table(
    ['Empresa', 'Utilização Média (%)', 'Colab. Atuais', 'Colab. Sugeridos', 'Diferença'],
    ocio_rows,
    caption='Tabela 10 – Top 10 empresas em ociosidade e potencial de redução de quadro (2026)'
)
blank()

txt('A coexistência de empresas em sobrecarga e em ociosidade dentro do mesmo portfólio configura uma ', B('ineficiência sistêmica de alocação de recursos humanos'), ': há capacidade ociosa disponível que não está sendo direcionada para onde a demanda é maior. Esse desequilíbrio é um dos principais subsídios que este trabalho oferece para o planejamento estratégico, sugerindo que mecanismos de realocação, compartilhamento de equipes ou ajuste de contratos poderiam equilibrar a carga entre as empresas, reduzindo custos de ociosidade sem comprometer o nível de serviço das empresas sobrecarregadas.')

# ──────────────────────────────────────────
h2('4.5 Síntese dos Resultados')
# ──────────────────────────────────────────

txt('O percurso analítico desenvolvido neste trabalho — desde a construção das bases de dados simuladas até a identificação de empresas críticas e ociosas — demonstra a viabilidade e a utilidade de um pipeline integrado de análise de capacidade produtiva e previsão de demanda implementado em Python.')

txt('A simulação com parâmetros controlados (sazonalidade definida, perfis de dimensionamento conhecidos, crescimento de 5% ao ano) permitiu validar cada etapa do processo: os modelos de previsão reproduziram os padrões sazonais esperados, os indicadores de capacidade diferenciaram corretamente os três perfis de empresa, e a análise mensal de 2026 identificou os períodos e empresas de maior risco operacional.')

txt('Os principais achados do estudo podem ser sintetizados nos seguintes pontos:')
txt('• A ', B('sazonalidade intra-anual'), ' é o fator determinante da pressão sobre a capacidade, com os meses de janeiro e dezembro representando os maiores riscos operacionais e exigindo planejamento antecipado de reforços de equipe;')
txt('• O modelo ', B('Prophet'), ' apresentou o melhor equilíbrio entre acurácia, interpretabilidade e facilidade de incorporação de domínio de negócio, sendo recomendado como ferramenta de previsão para horizontes de médio prazo neste tipo de ambiente;')
txt(f'• Aproximadamente ', B(f'{n_emp_crit} das 80 empresas'), ' operam sistematicamente em condição crítica (utilização > 90%), demandando reforço de quadro ou revisão dos processos para evitar degradação do serviço em 2026;')
txt(f'• Outros ', B(f'{n_emp_ocio} empresas'), ' apresentam ociosidade significativa (utilização < 70%), representando um potencial de otimização de custo de mão de obra que poderia ser realocado para as empresas sobrecarregadas;')
txt('• O indicador de ', B('colaboradores sugeridos'), ', derivado diretamente da demanda projetada e da jornada disponível, oferece uma referência objetiva e de fácil atualização para o dimensionamento de equipes em cada empresa e período do ano.')

txt('Esses resultados evidenciam que a integração entre previsão de demanda e análise de capacidade produtiva, implementada com dados granulares por empresa e produto, constitui uma ferramenta robusta para apoiar decisões estratégicas de planejamento e alocação de recursos humanos em ambientes de prestação de serviços, em linha com as diretrizes de Corrêa e Corrêa (2012) sobre gestão de capacidade efetiva.')

# ══════════════════════════════════════════════════════════════
# 5. ADICIONA REFERÊNCIAS FALTANTES
# ══════════════════════════════════════════════════════════════
# Encontra o último parágrafo de referências e adiciona após ele
last_ref = None
in_refs = False
for p in doc.paragraphs:
    if p.text.strip().startswith('Refer') and 'Heading' in p.style.name:
        in_refs = True
    if in_refs and p.text.strip():
        last_ref = p

if last_ref:
    def add_ref(text):
        p = doc.add_paragraph(style='Normal')
        p.add_run(text)
        last_ref._element.addnext(p._element)

    add_ref('BREIMAN, Leo. Random forests. Machine Learning, v. 45, n. 1, p. 5-32, 2001.')
    add_ref('TAYLOR, Sean J.; LETHAM, Benjamin. Forecasting at scale. The American Statistician, v. 72, n. 1, p. 37-45, 2018.')

# ══════════════════════════════════════════════════════════════
# 6. SALVA
# ══════════════════════════════════════════════════════════════
doc.save('TCC_Formatado.docx')
print("\n✓ TCC_Formatado.docx atualizado com sucesso!")
print(f"  Utilização média histórica  : {util_media_hist:.1f}%")
print(f"  Utilização média 2026       : {util_media_2026:.1f}%")
print(f"  Empresas críticas (2026)    : {n_emp_crit}")
print(f"  Empresas ociosas  (2026)    : {n_emp_ocio}")
print(f"  Total previsto Prophet 2026 : {total_prophet:,}")
